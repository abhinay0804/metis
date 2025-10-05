"""
DOCX content extractor using python-docx.
Extracts paragraphs and tables into a structured JSON.
"""

from typing import Dict, Any, List
from docx import Document


class DOCXExtractor:
    """Extract text, tables, and basic metadata from DOCX files."""

    def extract(self, file_path: str) -> Dict[str, Any]:
        content: Dict[str, Any] = {
            "paragraphs": [],
            "tables": [],
            "metadata": {}
        }
        try:
            doc = Document(file_path)

            # Core properties
            cp = doc.core_properties
            content["metadata"] = {
                "title": cp.title,
                "author": cp.author,
                "subject": cp.subject,
                "created": str(cp.created) if cp.created else None,
                "modified": str(cp.modified) if cp.modified else None,
            }

            # Paragraphs
            for p in doc.paragraphs:
                text = p.text.strip()
                if text:
                    content["paragraphs"].append(text)

            # Tables
            for table in doc.tables:
                table_rows: List[List[str]] = []
                for row in table.rows:
                    row_vals: List[str] = []
                    for cell in row.cells:
                        row_vals.append(cell.text.strip())
                    if any(v for v in row_vals):
                        table_rows.append(row_vals)
                if table_rows:
                    content["tables"].append({"data": table_rows})

            return content
        except Exception as e:
            raise Exception(f"Error extracting DOCX content: {str(e)}")


